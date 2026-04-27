"""Convert EDGE_CASES_TABLE.md to a formatted Word document."""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_rtl_para(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def parse_md_table(lines):
    """Parse markdown table lines into (headers, rows)."""
    table_lines = [l for l in lines if l.strip().startswith("|")]
    if len(table_lines) < 2:
        return [], []

    def split_row(line):
        parts = line.strip().strip("|").split("|")
        return [p.strip() for p in parts]

    headers = split_row(table_lines[0])
    rows = []
    for line in table_lines[2:]:  # skip separator line
        if re.match(r"^\s*\|[-| :]+\|\s*$", line):
            continue
        rows.append(split_row(line))
    return headers, rows


def strip_md(text: str) -> str:
    """Remove backticks and bold markers from cell text."""
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    return text.strip()


def main():
    with open("EDGE_CASES_TABLE.md", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Title
    title = doc.add_heading("EDGE_CASES_TABLE — מקרי קצה במערכת הנוכחית", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Split content into sections by "---" or "##"
    blocks = re.split(r"\n(?=#{1,3} |\-{3,})", content)

    current_section = ""
    for block in blocks:
        block = block.strip()
        if not block:
            continue

        # Section heading
        if block.startswith("## "):
            heading_text = block.split("\n")[0].lstrip("#").strip()
            h = doc.add_heading(heading_text, level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            current_section = heading_text
            # Process rest of block
            rest = "\n".join(block.split("\n")[1:]).strip()
            if rest:
                _process_block(doc, rest, current_section)
        elif block.startswith("### "):
            heading_text = block.split("\n")[0].lstrip("#").strip()
            h = doc.add_heading(heading_text, level=3)
            h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rest = "\n".join(block.split("\n")[1:]).strip()
            if rest:
                _process_block(doc, rest, current_section)
        elif block.startswith("---"):
            doc.add_paragraph()
        else:
            _process_block(doc, block, current_section)

    doc.save("EDGE_CASES_TABLE.docx")
    print("Saved: EDGE_CASES_TABLE.docx")


def _process_block(doc, block, section_name=""):
    lines = block.split("\n")

    # Check if block contains a markdown table
    table_lines = [l for l in lines if l.strip().startswith("|")]
    non_table_lines = [l for l in lines if not l.strip().startswith("|") and l.strip()]

    # Add non-table text as paragraphs
    for line in non_table_lines:
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(strip_md(line[2:]), style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p = doc.add_paragraph(strip_md(line))
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if not table_lines:
        return

    headers, rows = parse_md_table(table_lines)
    if not headers or not rows:
        return

    # Determine column widths based on section
    n_cols = len(headers)
    # Main edge case table has 8 columns
    if n_cols == 8:
        col_widths = [Cm(1.2), Cm(2.2), Cm(3.0), Cm(3.0), Cm(4.5), Cm(2.2), Cm(3.5), Cm(3.5)]
    elif n_cols == 6:
        col_widths = [Cm(1.5), Cm(2.5), Cm(4.0), Cm(4.0), Cm(3.5), Cm(3.5)]
    elif n_cols == 5:
        col_widths = [Cm(1.5), Cm(2.5), Cm(4.5), Cm(4.5), Cm(3.5)]
    else:
        col_widths = None

    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    if col_widths and len(col_widths) == n_cols:
        for i, cell in enumerate(table.columns[0].cells):
            pass  # will set per-row below

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, "1F3864")  # dark blue
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(strip_md(h))
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    STATUS_COLORS = {
        "קיים במערכת": "E2EFDA",      # light green
        "לא קיים במערכת": "FCE4D6",   # light red/orange
    }

    for row_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells

        # Determine row background from status column (index 5 for 8-col table)
        status_val = ""
        if n_cols == 8 and len(row_data) > 5:
            status_val = row_data[5].strip()
        elif n_cols >= 3 and len(row_data) > 2:
            status_val = row_data[2].strip()

        row_bg = "FFFFFF"
        if "קיים במערכת" in status_val and "לא" not in status_val:
            row_bg = "F2F9EE"
        elif "לא קיים" in status_val:
            row_bg = "FFF2EE"
        elif row_idx % 2 == 1:
            row_bg = "F5F5F5"

        for i, cell_text in enumerate(row_data):
            if i >= n_cols:
                break
            cell = row_cells[i]
            set_cell_bg(cell, row_bg)

            # Status column special coloring
            if n_cols == 8 and i == 5:
                if "קיים במערכת" in cell_text and "לא" not in cell_text:
                    set_cell_bg(cell, "C6EFCE")
                elif "לא קיים" in cell_text:
                    set_cell_bg(cell, "FFC7CE")

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl_para(p)
            run = p.add_run(strip_md(cell_text))
            run.font.size = Pt(7.5)

        # Set column widths
        if col_widths and len(col_widths) == n_cols:
            for i, w in enumerate(col_widths):
                if i < len(row_cells):
                    row_cells[i].width = w

    # Also set header row widths
    if col_widths and len(col_widths) == n_cols:
        for i, w in enumerate(col_widths):
            if i < len(hdr_cells):
                hdr_cells[i].width = w

    doc.add_paragraph()


if __name__ == "__main__":
    main()
