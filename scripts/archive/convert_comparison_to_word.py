"""Convert IMPLEMENTATION_COMPARISON.md to a formatted Word document."""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_rtl_para(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def strip_md(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    return text.strip()


def parse_md_table(lines):
    table_lines = [l for l in lines if l.strip().startswith("|")]
    if len(table_lines) < 2:
        return [], []

    def split_row(line):
        parts = line.strip().strip("|").split("|")
        return [p.strip() for p in parts]

    headers = split_row(table_lines[0])
    rows = []
    for line in table_lines[2:]:
        if re.match(r"^\s*\|[-| :]+\|\s*$", line):
            continue
        rows.append(split_row(line))
    return headers, rows


def add_table(doc, headers, rows, col_widths=None, header_bg="1F3864"):
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(strip_md(h))
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        row_bg = "F5F5F5" if row_idx % 2 == 1 else "FFFFFF"

        for i, cell_text in enumerate(row_data):
            if i >= n_cols:
                break
            cell = row_cells[i]

            # Special coloring for comparison columns
            if "לא קיים" in cell_text or "לא נבדק" in cell_text or "אין" == cell_text.strip() or "לא" == cell_text.strip():
                set_cell_bg(cell, "FFF2EE")
            elif "לא" in cell_text and ("מיושם" in cell_text or "רלוונטי" in cell_text):
                set_cell_bg(cell, "FFF2EE")
            elif cell_text.strip() in ("יש", "כן", "קיים"):
                set_cell_bg(cell, "F2F9EE")
            else:
                set_cell_bg(cell, row_bg)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl_para(p)
            run = p.add_run(strip_md(cell_text))
            run.font.size = Pt(7.5)

    # Set column widths
    if col_widths and len(col_widths) == n_cols:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = w

    doc.add_paragraph()
    return table


def main():
    with open("WEB_PATH_EDGE_CASES.md", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Page margins — landscape-friendly
    for section in doc.sections:
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        # Landscape
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)

    # Title
    title = doc.add_heading("השוואת מימוש: Web Path vs Direct-Excel Path", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if line.startswith("## "):
            h = doc.add_heading(line.lstrip("#").strip(), level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1

        elif line.startswith("### "):
            h = doc.add_heading(line.lstrip("#").strip(), level=3)
            h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1

        elif line.startswith("---"):
            doc.add_paragraph()
            i += 1

        elif line.startswith("|"):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            headers, rows = parse_md_table(table_lines)
            if not headers or not rows:
                continue

            n = len(headers)

            # Choose widths based on number of columns
            if n == 7:
                # Main comparison table
                widths = [Cm(2.5), Cm(3.5), Cm(4.5), Cm(3.5), Cm(4.5), Cm(3.5), Cm(3.0)]
            elif n == 4:
                widths = [Cm(1.5), Cm(4.0), Cm(8.0), Cm(8.0)]
            elif n == 3:
                widths = [Cm(3.5), Cm(8.0), Cm(8.0)]
            elif n == 2:
                widths = [Cm(8.0), Cm(12.0)]
            else:
                widths = None

            # Header color: dark blue for main table, teal for others
            hdr_color = "1F3864" if n == 7 else "2E4057"
            add_table(doc, headers, rows, col_widths=widths, header_bg=hdr_color)

        elif line and not line.startswith("#"):
            p = doc.add_paragraph(strip_md(line))
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1

        else:
            i += 1

    doc.save("WEB_PATH_EDGE_CASES.docx")
    print("Saved: WEB_PATH_EDGE_CASES.docx")


if __name__ == "__main__":
    main()
