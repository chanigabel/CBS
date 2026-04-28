"""Build UNIFIED_EDGE_CASES.docx — all edge cases with Web vs Direct-Excel side by side + examples."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_bg(cell, color):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), color)
    pr.append(s)


def set_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def write_cell(cell, text, size=7, bold=False, white=False):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


# ─────────────────────────────────────────────────────────────────────────────
# DATA TABLE
# Columns: (id, domain, edge_case, example_input, web_behavior, web_file_func,
#           excel_behavior, excel_file_func, shared_engine)
# ─────────────────────────────────────────────────────────────────────────────
ROWS = [
# SHEET / LOADING
("SH-01","SHEET","גיליון ריק","גיליון ללא שורות ותאים",
 "detect_table_region→None; SheetDataset(rows=[], skipped=True)",
 "ExcelToJsonExtractor.extract_sheet_to_json",
 "אותו מנגנון — גיליון מדולג",
 "ExcelReader.detect_table_region","ExcelReader"),

("SH-02","SHEET","אין כותרת מזוהה","גיליון עם נתונים ללא keywords מוכרים",
 "_score_header_row<3 → None → skipped=True",
 "ExcelReader.detect_table_region",
 "find_header מחזיר None → processor מדלג",
 "ExcelReader.find_header","ExcelReader"),

("SH-03","SHEET","כותרת אחרי שורה 30","כותרת בשורה 35",
 "max_scan_rows=30 → כותרת לא נמצאת → skipped",
 "ExcelToJsonExtractor.__init__",
 "לא רלוונטי — find_header סורק את כל הגיליון",
 "ExcelReader.find_header","—"),

("SH-04","SHEET","כותרות ממוזגות","A1:C1 ממוזגות עם 'שם פרטי'",
 "_is_merged_cell → קריאת ערך מהתא הימני-עליון",
 "ExcelReader._is_merged_cell",
 "_unmerge_header_area מבטל מיזוגים לפני הכנסת עמודות",
 "standardizationOrchestrator._unmerge_header_area","—"),

("SH-05","SHEET","שורת עזר מספרית זוהתה","שורה: 1,2,3,4,5 (רצופים, <100)",
 "_is_column_index_row=True → data_start_row+1",
 "ExcelReader._is_column_index_row",
 "_remove_numeric_helper_row מוחק שורה מהגיליון",
 "standardizationOrchestrator._remove_numeric_helper_row","—"),

("SH-06","SHEET","שורת עזר — לא רצופה","שורה: 1,3,5 (פערים)",
 "_is_column_index_row=False → שורה נשמרת",
 "ExcelReader._is_column_index_row",
 "לא מוחק — אותה בדיקה",
 "standardizationOrchestrator._remove_numeric_helper_row","—"),

("SH-07","SHEET","שורת עזר — פחות מ-3 ערכים","שורה: 1,2 (רק 2 תאים)",
 "len(values)<3 → False → שורה נשמרת",
 "ExcelReader._is_column_index_row",
 "לא מוחק",
 "standardizationOrchestrator._remove_numeric_helper_row","—"),

("SH-08","SHEET","קובץ xlsm","העלאת file.xlsm",
 "נשמר כ-.xlsm; ייצוא תמיד כ-.xlsx",
 "UploadService.handle_upload / ExportService.export",
 "keep_vba=True; פלט נשמר כ-.xlsm",
 "standardizationOrchestrator.process_workbook_json","—"),

# ROWS / DISPLAY
("RD-01","ROWS","שורה ריקה לחלוטין","כל תאי המקור None",
 "any(v not None)=False → מסוננת",
 "WorkbookService.get_sheet_data / ExportService.visible_rows",
 "לא מסוננת בעיבוד; מסוננת בייצוא ע\"י is_valid_data_row",
 "ExportEngine.is_valid_data_row","—"),

("RD-02","ROWS","שורה עם רווחים בלבד","כל תאים: '   '",
 "str(v).strip()='' → מסוננת",
 "WorkbookService.get_sheet_data",
 "לא מסוננת בעיבוד",
 "—","—"),

("RD-03","ROWS","ערכים מתוקנים בלבד, מקור ריק","first_name=None, first_name_corrected='יוסי'",
 "סינון בודק רק עמודות מקור → שורה מסוננת",
 "WorkbookService.get_sheet_data",
 "לא רלוונטי",
 "—","—"),

("RD-04","ROWS","שורה ראשונה מספרית מוסרת","שורה 1: 1,2,3,4,5",
 "_is_numeric_like כל ערך → clean_rows[1:]",
 "WorkbookService.get_sheet_data",
 "_remove_numeric_helper_row מוחק מהגיליון",
 "standardizationOrchestrator._remove_numeric_helper_row","—"),

("RD-05","ROWS","שורה שנייה מספרית לא מוסרת","שורה 2: 1,2,3",
 "בדיקה רק על clean_rows[0] → שורה 2 נשמרת",
 "WorkbookService.get_sheet_data",
 "לא מוחק שורה שנייה",
 "—","—"),

("RD-06","ROWS","סטטוס ללא עוגן","identifier_status קיים, id_number לא ב-field_names",
 "_anchor_to_status ריק → status בסוף display_columns",
 "WorkbookService.get_sheet_data",
 "לא רלוונטי",
 "—","—"),

("RD-07","ROWS","עמודה מתוקנת ללא מקור","first_name_corrected קיים, first_name לא ב-field_names",
 "מוצב בסוף display_columns דרך 'remaining keys'",
 "WorkbookService.get_sheet_data",
 "לא רלוונטי",
 "—","—"),

("RD-08","ROWS","_serial סינתטי","אין עמודת מספר סידורי",
 "detect_serial_field=None → _serial=1,2,3...",
 "derived_columns.apply_derived_columns",
 "לא קיים",
 "—","—"),

("RD-09","ROWS","MosadID חסר","אין תווית MosadID בגיליון",
 "scan_mosad_id=None → עמודה לא מוצגת",
 "WorkbookService.get_sheet_data / scan_mosad_id",
 "לא מוזרק",
 "—","—"),

("RD-10","ROWS","מספר תוויות MosadID","שתי תוויות בגיליון",
 "scan_mosad_id עוצר בהתאמה הראשונה",
 "scan_mosad_id",
 "לא רלוונטי",
 "—","—"),
]

ROWS += [
# NAMES
("NM-01","NAMES","None","first_name=None",
 "safe_to_string→'' → first_name_corrected=''",
 "TextProcessor.safe_to_string",
 "normalize_name(None)→''",
 "NameFieldProcessor._process_simple_name_field","TextProcessor"),

("NM-02","NAMES","מחרוזת ריקה","first_name=''",
 "clean_name('')→''",
 "TextProcessor.clean_name",
 "normalize_name('')→''",
 "NameFieldProcessor","TextProcessor"),

("NM-03","NAMES","רווחים בלבד","first_name='   '",
 "split()=[] → join→''",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-04","NAMES","תווי רוחב אפס","first_name='\\u200b\\u200c'",
 "_ZERO_WIDTH מסנן → ''",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-05","NAMES","ספרות בלבד","first_name='12345'",
 "שפה MIXED → ספרות נדחות → ''",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-06","NAMES","סמלים בלבד","first_name='@#$%'",
 "שפה MIXED → נדחים → ''",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-07","NAMES","ספרות + עברית","first_name='יוסי123'",
 "שפה HEBREW → ספרות נדחות → 'יוסי'",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-08","NAMES","ספרות + אנגלית","first_name='John123'",
 "שפה ENGLISH → ספרות נדחות → 'John'",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-09","NAMES","פיסוק בלבד","first_name='.,;:!?'",
 "שפה MIXED → נדחים → ''",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-10","NAMES","מקפים מגוונים","first_name='בן-דוד' / 'בן-דוד' (en-dash)",
 "_HYPHEN_CHARS → רווח → 'בן דוד'",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-11","NAMES","גרש/גרשיים","first_name='ז\"ל'",
 "גרשיים נדחים → 'זל' → remove_unwanted_tokens → ''",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-12","NAMES","פסיק/נקודה/סוגריים","first_name='(יוסי)'",
 "סוגריים נדחים → 'יוסי'",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-13","NAMES","ניקוד","first_name='יוסף' (עם ניקוד)",
 "ניקוד 1456-1479 נדחה → 'יסף'",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-14","NAMES","תואר בלבד","first_name='ד\"ר'",
 "→ 'דר' → remove_unwanted_tokens → ''",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-15","NAMES","ז\"ל בלבד","first_name='ז\"ל'",
 "→ 'זל' → remove_unwanted_tokens → ''",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-16","NAMES","עברית בלבד","first_name='יוסי'",
 "שפה HEBREW → 'יוסי'",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-17","NAMES","אנגלית בלבד","first_name='John'",
 "שפה ENGLISH → 'John'",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-18","NAMES","ערבוב עברית/אנגלית","first_name='יוסי John'",
 "שפה HEBREW (עברית>=אנגלית) → אנגלית נדחת → 'יוסי'",
 "TextProcessor.detect_language_dominance",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-19","NAMES","שוויון עברית=אנגלית","first_name='ab יב' (2+2)",
 "hebrew_count>=english_count → HEBREW → אנגלית נדחת",
 "TextProcessor.detect_language_dominance",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-20","NAMES","אין אותיות שורדות","first_name='123 @#$'",
 "שפה MIXED → הכל נדחה → ''",
 "TextProcessor.clean_name",
 "אותה לוגיקה",
 "NameFieldProcessor","TextProcessor"),

("NM-21","NAMES","שם פרטי = שם משפחה (טוקן אחד)","first_name='כהן', last_name='כהן'",
 "len(split)==1 → Stage A+B לא פועלים → 'כהן'",
 "NameEngine.remove_last_name_from_first_name",
 "אותה לוגיקה",
 "NameFieldProcessor._process_simple_name_field","NameEngine"),

("NM-22","NAMES","שם אב = שם משפחה (טוקן אחד)","father_name='כהן', last_name='כהן'",
 "Stage A: remove_substring→'' → מחזיר ''",
 "NameEngine.remove_last_name_from_father",
 "אותה לוגיקה",
 "NameFieldProcessor._process_father_name_field","NameEngine"),

("NM-23","NAMES","הסרה דרך Stage A","first_name='כהן יוסי', last_name='כהן'",
 "Stage A מסיר → 'יוסי'; Stage B לא רץ",
 "NameEngine.remove_last_name_from_first_name",
 "אותה לוגיקה",
 "NameFieldProcessor","NameEngine"),

("NM-24","NAMES","pattern נלמד מ-N שורות","5 שורות ראשונות עם שם משפחה בתחילת שם האב",
 "Web: 10 שורות; pattern מוחל על כל השורות",
 "standardizationPipeline.normalize_dataset",
 "5 שורות; pattern מוחל על כל השורות",
 "NameFieldProcessor.detect_father_name_pattern","NameEngine"),

("NM-25","NAMES","NONE בשם פרטי, REMOVE_FIRST בשם אב","<3 התאמות בשם פרטי, >=3 בשם אב",
 "שם פרטי: Stage B מדולג; שם אב: Stage B פועל",
 "standardizationPipeline / NameEngine",
 "אותה לוגיקה",
 "NameFieldProcessor","NameEngine"),
]

ROWS += [
# GENDER
("GN-01","GENDER","None","gender=None",
 "normalize_gender(None)→1 (זכר)",
 "GenderEngine.normalize_gender",
 "אותה לוגיקה",
 "GenderFieldProcessor.process_data","GenderEngine"),

("GN-02","GENDER","מחרוזת ריקה","gender=''",
 "value_str='' → ריק → 1",
 "GenderEngine.normalize_gender",
 "אותה לוגיקה",
 "GenderFieldProcessor.process_data","GenderEngine"),

("GN-03","GENDER","רווחים בלבד","gender='   '",
 "strip()='' → 1",
 "GenderEngine.normalize_gender",
 "אותה לוגיקה",
 "GenderFieldProcessor.process_data","GenderEngine"),

("GN-04","GENDER","ז","gender='ז'",
 "לא ב-FEMALE_PATTERNS → 1",
 "GenderEngine.normalize_gender",
 "אותה לוגיקה",
 "GenderFieldProcessor.process_data","GenderEngine"),

("GN-05","GENDER","נ","gender='נ'",
 "'נ' ב-FEMALE_PATTERNS → 2",
 "GenderEngine.normalize_gender",
 "אותה לוגיקה",
 "GenderFieldProcessor.process_data","GenderEngine"),

("GN-06","GENDER","male","gender='male'",
 "לא ב-FEMALE_PATTERNS → 1",
 "GenderEngine.normalize_gender",
 "אותה לוגיקה",
 "GenderFieldProcessor.process_data","GenderEngine"),

("GN-07","GENDER","female","gender='female'",
 "'female' ב-FEMALE_PATTERNS → 2",
 "GenderEngine.normalize_gender",
 "אותה לוגיקה",
 "GenderFieldProcessor.process_data","GenderEngine"),

("GN-08","GENDER","f / F","gender='f' או 'F'",
 "lower()→'f' ב-FEMALE_PATTERNS → 2",
 "GenderEngine.normalize_gender",
 "אותה לוגיקה",
 "GenderFieldProcessor.process_data","GenderEngine"),

("GN-09","GENDER","ערך לא מוכר","gender='unknown'",
 "לא ב-FEMALE_PATTERNS → 1 (ברירת מחדל)",
 "GenderEngine.normalize_gender",
 "אותה לוגיקה",
 "GenderFieldProcessor.process_data","GenderEngine"),

("GN-10","GENDER","מלכודת substring 'נ'","gender='נקבה'",
 "'נ' in 'נקבה'=True → 2",
 "GenderEngine.normalize_gender",
 "אותה לוגיקה",
 "GenderFieldProcessor.process_data","GenderEngine"),

("GN-11","GENDER","ערך משולב","gender='זכר/נקבה'",
 "'נ' in 'זכר/נקבה'=True → 2 (נקבה מנצחת)",
 "GenderEngine.normalize_gender",
 "אותה לוגיקה",
 "GenderFieldProcessor.process_data","GenderEngine"),

("GN-12","GENDER","ריק — pipeline","gender=None או ''",
 "pipeline: original is None/'' → gender_corrected=original (לא מפעיל engine)",
 "standardizationPipeline.apply_gender_standardization",
 "engine תמיד נקרא",
 "GenderFieldProcessor.process_data","—"),

("GN-13","GENDER","כותרות מרובות","שני עמודות 'מין' בגיליון",
 "מזהה עמודה אחת בלבד",
 "ExcelReader.detect_columns",
 "מעבד כל כותרת 'מין' (VBA FindAllHeaders)",
 "GenderFieldProcessor.find_headers","—"),
]

ROWS += [
# DATES
("DT-01","DATES","מפוצל תקין","birth_year=1990, month=5, day=15",
 "parse_from_split_columns → is_valid=True",
 "DateEngine.parse_from_split_columns",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-02","DATES","מפוצל — חלק חסר","birth_year=1990, month=None, day=15",
 "_has_split_date=False → parse_from_main_value(None) → 'תא ריק'",
 "DateEngine._has_split_date",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-03","DATES","יום לא תקין","birth_day=35",
 "_validate_date: dy>31 → 'יום לא תקין', is_valid=False",
 "DateEngine._validate_date",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-04","DATES","חודש לא תקין","birth_month=13",
 "_validate_date: mo>12 → 'חודש לא תקין'",
 "DateEngine._validate_date",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-05","DATES","תאריך בלתי אפשרי","birth_year=1990, month=2, day=30",
 "datetime(1990,2,30) → ValueError → 'תאריך לא קיים'",
 "DateEngine._validate_date",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-06","DATES","לא ניתן לפריקה","birth_year='abc'",
 "int(float('abc')) → exception → 'תוכן לא ניתן לפריקה'",
 "DateEngine.parse_from_split_columns",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-07","DATES","שנה בלבד (4 ספרות)","main_val='1990'",
 "1900<=1990<=2100 → year=1990, month=0, day=0, 'חסר חודש ויום'",
 "DateEngine._parse_numeric_date_string",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-08","DATES","4 ספרות — לא שנה","main_val='0312'",
 "int=312; לא בטווח → DMYY: d=0,m=3,yr=expand(12)",
 "DateEngine._parse_numeric_date_string",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-09","DATES","5 ספרות","main_val='12345'",
 "len!=4,6,8 → 'אורך תאריך לא תקין'",
 "DateEngine._parse_numeric_date_string",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-10","DATES","6 ספרות","main_val='150590'",
 "dy=15,mo=05,yr=expand(90) → _validate_date",
 "DateEngine._parse_numeric_date_string",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-11","DATES","7 ספרות","main_val='1234567'",
 "len!=4,6,8 → 'אורך תאריך לא תקין'",
 "DateEngine._parse_numeric_date_string",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-12","DATES","8 ספרות","main_val='15051990'",
 "dy=15,mo=05,yr=1990 → _validate_date",
 "DateEngine._parse_numeric_date_string",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-13","DATES","מספר סידורי Excel","raw_value=36526 (int)",
 "1<=36526<=2958465 → from_excel(36526) → תאריך",
 "DateEngine.parse_date_value",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-14","DATES","מספר סידורי אפס","raw_value=0",
 "1<=0=False → txt='0' → 'אורך תאריך לא תקין'",
 "DateEngine.parse_date_value",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-15","DATES","מספר סידורי גדול","raw_value=9999999",
 "9999999>2958465 → txt='9999999' → 'אורך תאריך לא תקין'",
 "DateEngine.parse_date_value",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-16","DATES","שם חודש אנגלי","main_val='15 January 2005'",
 "_contains_month_name=True → month=1,day=15,year=2005",
 "DateEngine._parse_mixed_month_numeric",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-17","DATES","שם חודש עברי","main_val='15 ינואר 2005'",
 "_extract_month_number('ינואר')=1 → month=1",
 "DateEngine._parse_mixed_month_numeric",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-18","DATES","ISO-like","main_val='1997-09-04T00:00:00'",
 "regex ^(d{4})-(d{2})-(d{2}) → yr=1997,mo=9,dy=4",
 "DateEngine.parse_date_value",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-19","DATES","קו נטוי","main_val='15/05/1990'",
 "'/' in txt → _parse_separated_date_string(DDMM)",
 "DateEngine._parse_separated_date_string",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-20","DATES","נקודה","main_val='15.05.1990'",
 "replace('.','/') → _parse_separated_date_string",
 "DateEngine.parse_date_value",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-21","DATES","טקסט לא מזוהה","main_val='לא תאריך'",
 "לא datetime/int/month/digits/ISO/slash → 'פורמט תאריך לא מזוהה'",
 "DateEngine.parse_date_value",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-22","DATES","יום/חודש ללא שנה","main_val='15/05'",
 "len(parts)==2 → שנה נוכחית מוזרקת",
 "DateEngine._parse_separated_date_string",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-23","DATES","שנה דו-ספרתית","main_val='15/05/90'",
 "yr=90<100 → _expand_two_digit_year(90)",
 "DateEngine._expand_two_digit_year",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-24","DATES","שנה לפני 1900","birth_year=1850",
 "validate_business_rules: year<1900 → 'שנה לפני 1900'",
 "DateEngine.validate_business_rules",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-25","DATES","תאריך לידה עתידי","birth_date > today",
 "date_val>today → 'תאריך לידה עתידי'",
 "DateEngine.validate_business_rules",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-26","DATES","תאריך כניסה עתידי","entry_date > today",
 "date_val>today → 'תאריך כניסה עתידי'",
 "DateEngine.validate_business_rules",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-27","DATES","גיל מעל 100","birth_year=1900",
 "age>100 → is_valid=True אך status='גיל מעל 100 (N שנים)'",
 "DateEngine.validate_business_rules",
 "אותה לוגיקה + צביעה צהובה בתא",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-28","DATES","כניסה ריקה — web","entry_date=None",
 "ENTRY_DATE + 'תא ריק' → status='', is_valid=False",
 "DateEngine.validate_business_rules",
 "לא רלוונטי (split columns)",
 "—","DateEngine"),

("DT-29","DATES","רכיבים לא תקינים נכתבים","birth_month=13",
 "_validate_date שומר year/month/day גם כשלא תקין; pipeline כותב ל-_corrected",
 "standardizationPipeline._normalize_date_field",
 "אותה לוגיקה",
 "DateFieldProcessor._process_date_field","DateEngine"),

("DT-30","DATES","DDMM hardcoded — web","כל תאריך מפוצל",
 "תמיד DateFormatPattern.DDMM",
 "standardizationPipeline._normalize_date_field",
 "מזהה אוטומטית DDMM/MMDD מהנתונים",
 "DateFieldProcessor.detect_date_format_pattern","—"),

("DT-31","DATES","זיהוי DDMM/MMDD — Excel בלבד","main_vals='30/01/1990'",
 "לא קיים",
 "—",
 "first>12 → DDMM; second>12 → MMDD",
 "DateFieldProcessor.detect_date_format_pattern","—"),

("DT-32","DATES","כניסה לפני לידה — Excel בלבד","entry=01/01/1990, birth=01/01/2000",
 "לא קיים",
 "—",
 "_validate_entry_vs_birth מוסיף אזהרה + ורוד",
 "standardizationOrchestrator._validate_entry_vs_birth","DateEngine"),

("DT-33","DATES","כניסה לפני לידה — web (לא מיושם)","entry=01/01/1990, birth=01/01/2000",
 "validate_entry_before_birth קיים ב-DateEngine אך לא נקרא מ-pipeline",
 "DateEngine.validate_entry_before_birth (לא מופעל)",
 "לא רלוונטי",
 "—","DateEngine"),

("DT-34","DATES","datetime object בעמודת שנה","year_col=datetime(1990,5,15)",
 "isinstance(year_val, datetime) → main_val_for_engine=year_val",
 "standardizationPipeline._normalize_date_field",
 "לא מטופל במפורש",
 "DateFieldProcessor._normalize_split_value","—"),
]

ROWS += [
# IDENTIFIERS
("ID-01","IDENTIFIERS","שניהם חסרים","id_number=None, passport=None",
 "early return לפני engine; id_number_corrected=None, passport_corrected=None",
 "standardizationPipeline.apply_identifier_standardization",
 "find_headers מחזיר False → processor מדלג",
 "IdentifierFieldProcessor.find_headers","—"),

("ID-02","IDENTIFIERS","דרכון בלבד","id_number=None, passport='AB123456'",
 "id_str='' → cleaned_passport='AB123456' → status='דרכון הוזן'",
 "IdentifierEngine.normalize_identifiers",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-03","IDENTIFIERS","sentinel 9999","id_number='9999'",
 "id_str=='9999' → id_str='' → מטופל כ'אין ת.ז.'",
 "IdentifierEngine.normalize_identifiers",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-04","IDENTIFIERS","ת.ז. עם אותיות","id_number='12A456789'",
 "'A' לא ספרה ולא dash → moved_to_passport=True",
 "IdentifierEngine._process_id_value",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-05","IDENTIFIERS","ת.ז. עם רווח","id_number='123 456789'",
 "' ' לא ספרה ולא dash → moved_to_passport=True",
 "IdentifierEngine._process_id_value",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-06","IDENTIFIERS","ת.ז. עם נקודה","id_number='123.456789'",
 "'.' לא ספרה ולא dash → moved_to_passport=True",
 "IdentifierEngine._process_id_value",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-07","IDENTIFIERS","ת.ז. עם מקף ASCII","id_number='123-456789'",
 "ord('-')=45 ב-DASH_CHARS; מקף מותר; ספרות: 9 → checksum",
 "IdentifierEngine._process_id_value",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-08","IDENTIFIERS","ת.ז. עם unicode dash","id_number='123\\u2013456789' (en-dash)",
 "ord(en-dash)=8211 ב-DASH_CHARS; מותר",
 "IdentifierEngine._process_id_value",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-09","IDENTIFIERS","ת.ז. קצרה (<4 ספרות)","id_number='123'",
 "digit_count<4 → moved_to_passport=True; reason='too_short'",
 "IdentifierEngine._process_id_value",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-10","IDENTIFIERS","ת.ז. ארוכה (>9 ספרות)","id_number='1234567890'",
 "digit_count>9 → moved_to_passport=True; reason='too_long'",
 "IdentifierEngine._process_id_value",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-11","IDENTIFIERS","כל אפסים","id_number='000000000'",
 "all(ch=='0') → return '', False, passport, False; status='ת.ז. לא תקינה'",
 "IdentifierEngine._process_id_value",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-12","IDENTIFIERS","אפס אחד","id_number='0'",
 "digits='0', digit_count=1<4 → moved_to_passport=True",
 "IdentifierEngine._process_id_value",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-13","IDENTIFIERS","כל ספרות זהות","id_number='111111111'",
 "len(set(padded))==1 → return '', False, passport, False; status='ת.ז. לא תקינה'",
 "IdentifierEngine._process_id_value",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-14","IDENTIFIERS","checksum תקין","id_number='039337423'",
 "validate_israeli_id=True; status='ת.ז. תקינה'; מחזיר ת.ז. מקורית",
 "IdentifierEngine.validate_israeli_id",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-15","IDENTIFIERS","checksum לא תקין","id_number='123456789'",
 "validate_israeli_id=False; status='ת.ז. לא תקינה'; מחזיר cleaned_digits",
 "IdentifierEngine.validate_israeli_id",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-16","IDENTIFIERS","ת.ז. תקינה + דרכון","id_number='039337423', passport='AB123'",
 "checksum תקין; status='ת.ז. תקינה + דרכון הוזן'",
 "IdentifierEngine.normalize_identifiers",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-17","IDENTIFIERS","ת.ז. לא תקינה + דרכון","id_number='123456789', passport='AB123'",
 "checksum לא תקין; status='ת.ז. לא תקינה + דרכון הוזן'",
 "IdentifierEngine.normalize_identifiers",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-18","IDENTIFIERS","ריפוד 4 ספרות","id_number='1234'",
 "pad_id('1234')→'000001234'; checksum על '000001234'",
 "IdentifierEngine.pad_id",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-19","IDENTIFIERS","ריפוד 8 ספרות","id_number='12345678'",
 "pad_id('12345678')→'012345678'; checksum",
 "IdentifierEngine.pad_id",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-20","IDENTIFIERS","ת.ז. float מ-Excel","id_number=123456789.0",
 "_safe_to_string→'123456789.0'; '.' לא ספרה → moved_to_passport=True",
 "IdentifierEngine._process_id_value",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-21","IDENTIFIERS","דרכון עם רווחים","passport='AB 123 456'",
 "clean_passport: רווח נדחה → 'AB123456'",
 "IdentifierEngine.clean_passport",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-22","IDENTIFIERS","דרכון עם פיסוק","passport='AB-123.456'",
 "'.' נדחה; '-' (ord=45) ב-DASH_CHARS → נשמר → 'AB-123456'",
 "IdentifierEngine.clean_passport",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-23","IDENTIFIERS","דרכון עם אותיות עבריות","passport='אב123'",
 "1488<=ord('א')<=1514 → נשמר → 'אב123'",
 "IdentifierEngine.clean_passport",
 "אותה לוגיקה",
 "IdentifierFieldProcessor.process_data","IdentifierEngine"),

("ID-24","IDENTIFIERS","דרישת שתי כותרות — Excel","רק עמודת ת.ז., ללא דרכון",
 "לא רלוונטי — web מעבד גם אם רק אחד קיים",
 "standardizationPipeline.apply_identifier_standardization",
 "find_headers מחזיר False → processor מדלג לחלוטין",
 "IdentifierFieldProcessor.find_headers","—"),

("ID-25","IDENTIFIERS","שורה ללא שדות מזהה","אין id_number ואין passport",
 "early return ללא שינוי",
 "standardizationPipeline.apply_identifier_standardization",
 "לא רלוונטי",
 "—","—"),
]

ROWS += [
# EDIT / DELETE / RE-NORMALIZE
("ED-01","EDIT","עריכת שדה מקורי","PATCH /cell עם field_name='first_name'",
 "מעדכן dict בזיכרון + record.edits; first_name_corrected לא מתעדכן",
 "EditService.edit_cell",
 "לא קיים",
 "—","—"),

("ED-02","EDIT","עריכת שדה מתוקן","PATCH /cell עם field_name='first_name_corrected'",
 "מעדכן ישירות אם המפתח קיים בשורה",
 "EditService.edit_cell",
 "לא קיים",
 "—","—"),

("ED-03","EDIT","אינדקס שורה לא תקין","row_index=-1 או >=len(rows)",
 "HTTP 400",
 "EditService.edit_cell",
 "לא קיים",
 "—","—"),

("ED-04","EDIT","שם שדה לא תקין","field_name='nonexistent'",
 "HTTP 400 עם רשימת שדות זמינים",
 "EditService.edit_cell",
 "לא קיים",
 "—","—"),

("ED-05","EDIT","מחיקה עם אינדקסים כפולים","row_indices=[2,2,5]",
 "unique_indices=sorted(set([2,2,5]))=[2,5]; מחיקת 2 שורות",
 "EditService.delete_rows",
 "לא קיים",
 "—","—"),

("ED-06","EDIT","מחיקה עם אינדקס לא תקין","row_indices=[1,999]",
 "invalid=[999] → HTTP 400; אף שורה לא נמחקת (all-or-nothing)",
 "EditService.delete_rows",
 "לא קיים",
 "—","—"),

("ED-07","EDIT","מחיקת כל השורות","row_indices=[0..n-1]",
 "sheet.rows=[] → remaining_rows=0",
 "EditService.delete_rows",
 "לא קיים",
 "—","—"),

("ED-08","EDIT","מחיקה ואז re-normalize","מחיקת שורות → POST /normalize",
 "normalize מחלץ מחדש מהקובץ → שורות שנמחקו חוזרות",
 "standardizationService.normalize",
 "לא רלוונטי",
 "—","—"),

("ED-09","EDIT","re-normalize מבטל עריכות","עריכת תא → POST /normalize",
 "normalize מחלץ מחדש; record.edits לא מוחל מחדש → עריכות אובדות",
 "standardizationService.normalize",
 "לא רלוונטי",
 "—","—"),

("ED-10","EDIT","edits dict לא מוחל מחדש","record.edits מכיל עריכות",
 "edits נשמר ב-SessionRecord אך אין קוד שמחיל אותו לאחר normalize",
 "SessionRecord.edits / standardizationService.normalize",
 "לא רלוונטי",
 "—","—"),

("ED-11","EDIT","normalize כל הגיליונות","POST /normalize ללא ?sheet=",
 "sheet_name=None → כל הגיליונות מחולצים ומנורמלים",
 "standardizationService.normalize",
 "לא רלוונטי",
 "—","—"),

("ED-12","EDIT","normalize גיליון אחד","POST /normalize?sheet=שם",
 "רק הגיליון המבוקש מחולץ ומנורמל; שאר לא נגעים",
 "standardizationService.normalize",
 "לא רלוונטי",
 "—","—"),
]

ROWS += [
# EXPORT
("EX-01","EXPORT","ייצוא לפני נורמליזציה","POST /export ללא POST /normalize",
 "workbook_dataset is None → מחלץ מהקובץ; ייצוא ערכים מקוריים ללא תיקון",
 "ExportService.export",
 "לא רלוונטי",
 "—","—"),

("EX-02","EXPORT","שדות מתוקנים חסרים","שורה ללא first_name_corrected",
 "_cell_value→None; תא ריק; אין fallback לשדה מקורי",
 "ExportService._cell_value / EXPORT_MAPPING",
 "pick(corrected, original) — יש fallback",
 "ExportEngine._map_row_to_export_fields","—"),

("EX-03","EXPORT","קלט xlsm — פלט xlsx","קובץ .xlsm מועלה",
 "ExportService יוצר Workbook() חדש; שומר כ-.xlsx",
 "ExportService.export",
 "שומר .xlsm לקבצי מאקרו",
 "standardizationOrchestrator.process_workbook_json","—"),

("EX-04","EXPORT","אין הדגשה בייצוא","ייצוא web path",
 "Workbook() חדש; אין העתקת עיצוב; אין צביעה ורודה",
 "ExportService.export",
 "ורוד/צהוב בתאים שהשתנו",
 "ExcelWriter.highlight_changed_cells / format_cell","—"),

("EX-05","EXPORT","גיליון RTL","כל גיליון בייצוא",
 "ws.sheet_view.rightToLeft=True",
 "ExportService.export",
 "לא — שומר כיוון מקורי",
 "—","—"),

("EX-06","EXPORT","SugMosad לא מאוכלס","SugMosad ב-EXPORT_MAPPING",
 "_cell_value(row,'SugMosad'): אם ריק → None; תא ריק; לא מחושב",
 "ExportService.EXPORT_MAPPING",
 "לא רלוונטי",
 "—","—"),

("EX-07","EXPORT","שורה מוחרגת — מקור ריק","כל עמודות המקור ריקות",
 "visible_rows מסנן; שורה לא מופיעה",
 "ExportService.visible_rows",
 "is_valid_data_row: ANY שדה אישי לא ריק",
 "ExportEngine.is_valid_data_row","—"),

("EX-08","EXPORT","שורה נכללת — שדה אחד לא ריק","first_name='יוסי', שאר ריקים",
 "any(v not None)=True → שורה נכללת",
 "ExportService.visible_rows",
 "is_valid_data_row=True",
 "ExportEngine.is_valid_data_row","—"),

("EX-09","EXPORT","שורה שנמחקה לא מופיעה","מחיקת שורה → ייצוא",
 "שורה נמחקת מ-sheet.rows; visible_rows לא רואה אותה",
 "ExportService.visible_rows / EditService.delete_rows",
 "לא רלוונטי",
 "—","—"),

("EX-10","EXPORT","שורה שנמחקה חוזרת לאחר re-normalize","מחיקה → normalize → export",
 "normalize מחלץ מחדש מהקובץ; שורה חוזרת; ייצוא כולל אותה",
 "standardizationService.normalize / ExportService.export",
 "לא רלוונטי",
 "—","—"),

("EX-11","EXPORT","שם גיליון לא מוכר — fallback schema","שם גיליון לא תואם אף pattern",
 "canonical_sheet_name מחזיר שם מקורי; headers_for_sheet→_HEADERS_DEFAULT",
 "ExportService.canonical_sheet_name / headers_for_sheet",
 "לא רלוונטי",
 "—","—"),

("EX-12","EXPORT","bulk export — session נכשל","session אחד מתוך רשימה נכשל",
 "warning + skip; אם כולם נכשלו → HTTP 500",
 "export.export_bulk",
 "לא רלוונטי",
 "—","—"),

("EX-13","EXPORT","bulk export ריק","session_ids=[]",
 "HTTP 400",
 "export.export_bulk",
 "לא רלוונטי",
 "—","—"),
]


# ─────────────────────────────────────────────────────────────────────────────
# BUILD WORD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

HEADERS = [
    "מזהה", "תחום", "מקרה קצה", "קלט לדוגמה",
    "Web — מה קורה", "Web — קובץ/פונקציה",
    "Direct-Excel — מה קורה", "Direct-Excel — קובץ/פונקציה",
    "מנוע משותף"
]

# Column widths in cm (landscape A4 = ~25.4cm usable)
COL_WIDTHS = [1.2, 1.8, 2.8, 3.2, 4.2, 3.5, 4.2, 3.5, 2.0]

# Domain → background color for data rows
DOMAIN_COLORS = {
    "SHEET":       "EBF3FB",
    "ROWS":        "FFF9E6",
    "NAMES":       "F0FBF0",
    "GENDER":      "FDF0FB",
    "DATES":       "FFF0F0",
    "IDENTIFIERS": "F0F8FF",
    "EDIT":        "FFFAF0",
    "EXPORT":      "F5F5F5",
}

# Status-based override colors
WEB_ONLY_COLOR   = "E8F5E9"   # green tint — web only
EXCEL_ONLY_COLOR = "FFF3E0"   # orange tint — excel only
NOT_IMPL_COLOR   = "FFEBEE"   # red tint — not implemented


def domain_section_header(doc, domain_name):
    """Add a shaded section heading row spanning all columns."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"  {domain_name}  ")
    run.bold = True
    run.font.size = Pt(9)


def build_doc():
    doc = Document()

    # Landscape A4
    for section in doc.sections:
        section.page_width  = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin  = Cm(1.0)
        section.right_margin = Cm(1.0)
        section.top_margin   = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # Title
    title = doc.add_heading("טבלת מקרי קצה מאוחדת — Web vs Direct-Excel", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    subtitle = doc.add_paragraph("כל מקרה קצה בשורה אחת | Web ו-Direct-Excel זה לצד זה | כולל דוגמאות")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Create table
    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(HEADERS):
        set_bg(hdr[i], "1F3864")
        write_cell(hdr[i], h, size=7.5, bold=True, white=True)

    set_col_widths(table, COL_WIDTHS)

    # Data rows
    prev_domain = None
    for row_data in ROWS:
        (rid, domain, edge_case, example,
         web_what, web_where,
         excel_what, excel_where,
         engine) = row_data

        # Domain separator row
        if domain != prev_domain:
            sep_row = table.add_row().cells
            for c in sep_row:
                set_bg(c, "2E4057")
            # Merge all cells for the separator
            sep_row[0].merge(sep_row[-1])
            write_cell(sep_row[0], f"  {domain}  ", size=8, bold=True, white=True)
            prev_domain = domain

        tr = table.add_row().cells

        # Row background
        base_color = DOMAIN_COLORS.get(domain, "FFFFFF")

        # Detect special cases
        web_na    = web_what in ("לא קיים", "—", "לא רלוונטי") or web_where == "—"
        excel_na  = excel_what in ("לא קיים", "—", "לא רלוונטי") or excel_where == "—"

        for i, cell in enumerate(tr):
            set_bg(cell, base_color)

        # Override colors for web/excel cells
        if web_na and not excel_na:
            set_bg(tr[4], EXCEL_ONLY_COLOR)
            set_bg(tr[5], EXCEL_ONLY_COLOR)
        elif excel_na and not web_na:
            set_bg(tr[6], WEB_ONLY_COLOR)
            set_bg(tr[7], WEB_ONLY_COLOR)

        # "לא מיושם" highlight
        if "לא מיושם" in web_what or "לא נקרא" in web_what:
            set_bg(tr[4], NOT_IMPL_COLOR)
        if "לא מיושם" in excel_what or "לא נקרא" in excel_what:
            set_bg(tr[6], NOT_IMPL_COLOR)

        write_cell(tr[0], rid,        size=7, bold=True)
        write_cell(tr[1], domain,     size=7)
        write_cell(tr[2], edge_case,  size=7, bold=True)
        write_cell(tr[3], example,    size=6.5)
        write_cell(tr[4], web_what,   size=6.5)
        write_cell(tr[5], web_where,  size=6)
        write_cell(tr[6], excel_what, size=6.5)
        write_cell(tr[7], excel_where,size=6)
        write_cell(tr[8], engine,     size=6.5)

    set_col_widths(table, COL_WIDTHS)

    # Legend
    doc.add_paragraph()
    legend = doc.add_paragraph()
    legend.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    legend.add_run("מקרא: ").bold = True
    legend.add_run("ירוק = קיים רק ב-Web  |  כתום = קיים רק ב-Direct-Excel  |  אדום = קיים בקוד אך לא פעיל")

    out = "UNIFIED_EDGE_CASES.docx"
    doc.save(out)
    print(f"Saved: {out}  ({len(ROWS)} rows)")


if __name__ == "__main__":
    build_doc()
